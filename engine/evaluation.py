"""
evaluation.py — additive evaluation and submission-matrix helpers for hackathon artifacts.

This module keeps the existing scoring engine intact while adding:
- streaming JSONL candidate ingestion
- lightweight schema validation
- trap/anomaly detection for malformed candidate payloads
- submission-matrix normalization for CSV export conformance
- docx-driven rule/spec extraction when the official files are present
"""
from __future__ import annotations

import csv
import json
import logging
import re
from pathlib import Path
from typing import Any, Iterator, List, Optional, Sequence

import pandas as pd

logger = logging.getLogger(__name__)

BASE_DIR = Path(__file__).resolve().parent.parent


HACKATHON_RULE_SYNONYMS = {
    "schema_integrity": [
        "candidate schema",
        "schema validation",
        "required keys",
        "required fields",
    ],
    "chronology_integrity": [
        "chronology integrity",
        "chronological order",
        "timeline consistency",
        "date order",
        "inverted timeline",
    ],
    "duplicate_skill_trap": [
        "duplicate skill list",
        "uniform skill list",
        "repeated skills",
        "skill spam",
        "skill trick",
    ],
    "submission_matrix": [
        "sample submission",
        "submission matrix",
        "validate submission",
        "csv export",
        "tracking format",
    ],
}


def resolve_artifact_path(filename: str) -> Optional[Path]:
    """Find a hackathon artifact in the repo root or common nearby locations."""
    candidate_paths = [
        Path.cwd() / filename,
        BASE_DIR / filename,
        BASE_DIR / "data" / filename,
    ]
    for candidate in candidate_paths:
        if candidate.exists():
            return candidate
    return None


def load_json_schema(schema_path: str = "candidate_schema.json") -> dict:
    """Load the official candidate schema if present; otherwise return a permissive default."""
    resolved = resolve_artifact_path(schema_path)
    if not resolved:
        return {"type": "object", "required": [], "properties": {}}
    try:
        with open(resolved, "r", encoding="utf-8") as schema_file:
            return json.load(schema_file)
    except Exception as exc:
        logger.warning("Failed to load candidate schema from %s: %s", resolved, exc)
        return {"type": "object", "required": [], "properties": {}}


def stream_jsonl_candidates(candidate_path: str = "candidates.jsonl") -> Iterator[dict]:
    """Yield candidate records from a JSONL file without loading the full file into memory."""
    resolved = resolve_artifact_path(candidate_path)
    if resolved is None:
        logger.info("Candidate JSONL file not found: %s", candidate_path)
        return

    path_to_open = str(resolved)
    if resolved.name == "candidates.jsonl" and resolved == Path.cwd() / "candidates.jsonl":
        path_to_open = "candidates.jsonl"

    with open(path_to_open, "r", encoding="utf-8") as f:
        for line_number, line in enumerate(f, start=1):
            raw = line.strip()
            if not raw:
                continue
            try:
                record = json.loads(raw)
            except Exception as exc:
                logger.warning("Skipping malformed JSONL row %s in %s: %s", line_number, path_to_open, exc)
                continue
            if isinstance(record, dict):
                yield record
            else:
                logger.warning("Skipping non-object JSONL row %s in %s", line_number, path_to_open)


def _type_matches(value: Any, schema_type: Any) -> bool:
    if schema_type is None:
        return True
    if isinstance(schema_type, list):
        return any(_type_matches(value, item) for item in schema_type)
    if schema_type == "object":
        return isinstance(value, dict)
    if schema_type == "array":
        return isinstance(value, list)
    if schema_type == "string":
        return isinstance(value, str)
    if schema_type == "number":
        return isinstance(value, (int, float)) and not isinstance(value, bool)
    if schema_type == "integer":
        return isinstance(value, int) and not isinstance(value, bool)
    if schema_type == "boolean":
        return isinstance(value, bool)
    return True


def validate_candidate_record(record: dict, schema: Optional[dict] = None) -> List[str]:
    """Validate a candidate record against a JSON-schema-like structure."""
    errors: List[str] = []
    if not isinstance(record, dict):
        return ["Record is not a dictionary"]

    schema = schema or {"type": "object", "required": [], "properties": {}}
    required = schema.get("required", []) or []
    properties = schema.get("properties", {}) or {}

    for key in required:
        if key not in record:
            errors.append(f"Missing required key: {key}")

    for key, spec in properties.items():
        if key not in record:
            continue
        value = record[key]
        expected_type = spec.get("type") if isinstance(spec, dict) else None
        if expected_type and not _type_matches(value, expected_type):
            errors.append(f"Key '{key}' expected type {expected_type!r}")
            continue

        if isinstance(spec, dict) and expected_type == "array" and "items" in spec and isinstance(value, list):
            item_schema = spec.get("items", {})
            item_type = item_schema.get("type") if isinstance(item_schema, dict) else None
            for idx, item in enumerate(value):
                if item_type and not _type_matches(item, item_type):
                    errors.append(f"Key '{key}' item {idx} expected type {item_type!r}")

        if isinstance(spec, dict) and expected_type == "object" and isinstance(value, dict):
            nested_required = spec.get("required", []) or []
            nested_properties = spec.get("properties", {}) or {}
            for nested_key in nested_required:
                if nested_key not in value:
                    errors.append(f"Key '{key}' missing nested key: {nested_key}")
            for nested_key, nested_spec in nested_properties.items():
                if nested_key not in value:
                    continue
                nested_value = value[nested_key]
                nested_type = nested_spec.get("type") if isinstance(nested_spec, dict) else None
                if nested_type and not _type_matches(nested_value, nested_type):
                    errors.append(f"Key '{key}.{nested_key}' expected type {nested_type!r}")

    return errors


def detect_trap_anomalies(record: dict) -> List[str]:
    """Spot malformed or adversarial candidate payloads without changing score computation."""
    warnings: List[str] = []
    if not isinstance(record, dict):
        return ["Record is not a dictionary"]

    lowered_keys = {str(key).lower(): key for key in record.keys()}

    def _coerce_year(value: Any) -> Optional[int]:
        if isinstance(value, int) and not isinstance(value, bool):
            return value
        if isinstance(value, str):
            match = re.search(r"\b(19\d{2}|20[0-3]\d)\b", value)
            if match:
                return int(match.group(1))
        return None

    start_keys = [key for key in lowered_keys if any(token in key for token in ("start", "from", "join", "begin"))]
    end_keys = [key for key in lowered_keys if any(token in key for token in ("end", "to", "current", "present", "finish"))]
    if start_keys and end_keys:
        start_value = _coerce_year(record[lowered_keys[start_keys[0]]])
        end_value = _coerce_year(record[lowered_keys[end_keys[0]]])
        if start_value is not None and end_value is not None and start_value > end_value:
            warnings.append("Chronology appears inverted")

    skill_keys = [key for key in lowered_keys if "skill" in key]
    for skill_key in skill_keys:
        skill_value = record[lowered_keys[skill_key]]
        if isinstance(skill_value, list) and skill_value:
            normalized = [str(item).strip().lower() for item in skill_value if str(item).strip()]
            if len(set(normalized)) <= 1:
                warnings.append(f"Uniform or duplicated skill list detected in '{skill_key}'")
            elif len(set(normalized)) / max(len(normalized), 1) < 0.5:
                warnings.append(f"High duplicate ratio in '{skill_key}'")

    experience_value = record.get("years_of_experience") or record.get("experience_years") or record.get("years")
    if isinstance(experience_value, (int, float)) and not isinstance(experience_value, bool):
        if experience_value < 0:
            warnings.append("Negative experience value detected")
        elif experience_value > 50:
            warnings.append("Unusually large experience value detected")

    return warnings


def _coerce_text(record: dict) -> str:
    for key in ("resume_text", "text", "candidate_text", "profile_text", "description", "summary"):
        value = record.get(key)
        if isinstance(value, str) and value.strip():
            return value
    return json.dumps(record, ensure_ascii=False, sort_keys=True)


def _coerce_name(record: dict, index: Optional[int] = None) -> str:
    for key in ("name", "full_name", "candidate_name", "applicant_name"):
        value = record.get(key)
        if isinstance(value, str) and value.strip():
            return value.strip()
    return f"Candidate {index if index is not None else 'Unknown'}"


def _coerce_contact(record: dict) -> dict:
    contact = record.get("contact") if isinstance(record.get("contact"), dict) else {}
    email = record.get("email") or contact.get("email") or "Not Provided"
    phone = record.get("phone") or contact.get("phone") or "Not Provided"
    return {"email": email, "phone": phone}


def normalize_candidate_record(record: dict, index: Optional[int] = None) -> Optional[dict]:
    """Convert a raw candidate payload into the app's scoring shape."""
    if not isinstance(record, dict):
        return None

    name = _coerce_name(record, index=index)
    text = _coerce_text(record)
    normalized = {
        "name": name,
        "filename": record.get("filename") or record.get("id") or f"candidate_{index if index is not None else 'unknown'}",
        "text": text,
        "contact": _coerce_contact(record),
        "raw_record": record,
        "source": record.get("source") or "hackathon_dataset",
    }
    return normalized


def load_validated_candidates(
    candidate_path: str = "candidates.jsonl",
    schema_path: str = "candidate_schema.json",
) -> List[dict]:
    """Materialize only validated candidates, logging malformed rows instead of failing."""
    schema = load_json_schema(schema_path)
    validated: List[dict] = []
    for index, record in enumerate(stream_jsonl_candidates(candidate_path), start=1):
        schema_errors = validate_candidate_record(record, schema)
        trap_warnings = detect_trap_anomalies(record)
        if schema_errors:
            logger.warning("Skipping invalid candidate row %s: %s", index, "; ".join(schema_errors))
            continue
        if trap_warnings:
            logger.warning("Candidate row %s has potential anomalies: %s", index, "; ".join(trap_warnings))
        normalized = normalize_candidate_record(record, index=index)
        if normalized:
            validated.append(normalized)
    return validated


def extract_docx_spec_lines(docx_path: str) -> List[str]:
    """Extract ordered lines from a spec docx if the file exists."""
    resolved = resolve_artifact_path(docx_path)
    if not resolved:
        return []

    try:
        from docx import Document
    except Exception:
        logger.warning("python-docx is not available; skipping %s", resolved)
        return []

    try:
        document = Document(str(resolved))
    except Exception as exc:
        logger.warning("Failed to read docx spec %s: %s", resolved, exc)
        return []

    lines: List[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    lines.append(text)
    return lines


def extract_metric_rules_from_docs(docx_paths: Sequence[str] = ("redrob_signals_doc.docx", "submission_spec.docx")) -> dict:
    """Pull validation rules from the official docs when they are present."""
    collected_lines: List[str] = []
    for docx_path in docx_paths:
        collected_lines.extend(extract_docx_spec_lines(docx_path))

    rule_lines = [
        line for line in collected_lines
        if re.search(r"\b(score|submission|schema|candidate|export|validate|trap|duplicate|chronolog)\b", line, re.IGNORECASE)
    ]
    return {
        "source_docs": list(docx_paths),
        "rule_lines": rule_lines,
        "synonyms": HACKATHON_RULE_SYNONYMS,
    }


def _load_submission_columns(spec_path: str = "sample_submission.csv", validator_path: str = "validate_submission.py") -> List[str]:
    resolved_spec = resolve_artifact_path(spec_path)
    if resolved_spec and resolved_spec.exists():
        try:
            with open(resolved_spec, "r", encoding="utf-8", newline="") as csv_file:
                reader = csv.reader(csv_file)
                header = next(reader, [])
                if header:
                    return list(header)
        except Exception as exc:
            logger.warning("Failed to read sample submission columns from %s: %s", resolved_spec, exc)

    resolved_validator = resolve_artifact_path(validator_path)
    if resolved_validator and resolved_validator.exists():
        try:
            content = resolved_validator.read_text(encoding="utf-8")
            matches = re.findall(r"[\"']([A-Za-z0-9_ %()\-/%]+)[\"']", content)
            columns = [match for match in matches if any(token in match.lower() for token in ("score", "name", "rank", "email", "job", "status", "experience", "education", "semantic", "phone"))]
            if columns:
                seen = set()
                ordered_columns: List[str] = []
                for column in columns:
                    if column not in seen:
                        ordered_columns.append(column)
                        seen.add(column)
                return ordered_columns
        except Exception as exc:
            logger.warning("Failed to infer submission columns from %s: %s", resolved_validator, exc)

    return [
        "Rank",
        "Name",
        "Final Score (%)",
        "Matched Role",
        "Skill Match",
        "Experience Match",
        "Education Match",
        "Semantic Similarity",
        "Years Experience",
        "Status",
        "Email",
    ]


def _format_percentage(value: Any) -> str:
    if value is None:
        return "0.0%"
    if isinstance(value, str):
        return value if value.endswith("%") else f"{value}%"
    try:
        return f"{float(value):.1f}%"
    except Exception:
        return "0.0%"


def build_submission_matrix_rows(candidates: Sequence[dict], spec_path: str = "sample_submission.csv") -> List[dict]:
    """Build CSV-ready rows that follow the sample submission column order."""
    columns = _load_submission_columns(spec_path=spec_path)
    rows: List[dict] = []

    for default_rank, candidate in enumerate(candidates, start=1):
        breakdown = candidate.get("breakdown") or {}
        row = {
            "Rank": candidate.get("rank", default_rank),
            "Name": candidate.get("name", "Unknown"),
            "Final Score (%)": _format_percentage(candidate.get("final_score", 0.0)),
            "Matched Role": candidate.get("best_fit_job_title", candidate.get("matched_role", "None")),
            "Skill Match": _format_percentage(breakdown.get("skill_score", 0.0)),
            "Experience Match": _format_percentage(breakdown.get("experience_score", 0.0)),
            "Education Match": _format_percentage(breakdown.get("education_score", 0.0)),
            "Semantic Similarity": _format_percentage(breakdown.get("semantic_score", 0.0)),
            "Years Experience": candidate.get("candidate_years", 0.0),
            "Status": candidate.get("status", ""),
            "Email": candidate.get("contact", {}).get("email", "Not Provided"),
        }

        for column in columns:
            row.setdefault(column, row.get(column, ""))
        rows.append({column: row.get(column, "") for column in columns})

    return rows


def build_submission_matrix(candidates: Sequence[dict], spec_path: str = "sample_submission.csv") -> pd.DataFrame:
    rows = build_submission_matrix_rows(candidates, spec_path=spec_path)
    columns = _load_submission_columns(spec_path=spec_path)
    if not rows:
        return pd.DataFrame(columns=columns)
    return pd.DataFrame(rows, columns=columns)