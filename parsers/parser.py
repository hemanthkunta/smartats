"""
parser.py — extracts clean raw text from PDF / DOCX / TXT resumes and job descriptions.
"""
import re
import io
import PyPDF2
import docx


def extract_text_from_pdf(file_bytes: bytes) -> str:
    text_chunks = []
    reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
    for page in reader.pages:
        page_text = page.extract_text() or ""
        text_chunks.append(page_text)
    return "\n".join(text_chunks)


def extract_text_from_docx(file_bytes: bytes) -> str:
    document = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in document.paragraphs]
    # also grab text inside tables (lots of resumes use table layouts)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.append(cell.text)
    return "\n".join(paragraphs)


def extract_text_from_txt(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="ignore")


def extract_text(filename: str, file_bytes: bytes) -> str:
    """Dispatch to the right extractor based on file extension with error resilience."""
    if not file_bytes:
        return ""
    
    lower = filename.lower()
    try:
        if lower.endswith(".pdf"):
            raw = extract_text_from_pdf(file_bytes)
        elif lower.endswith(".docx"):
            raw = extract_text_from_docx(file_bytes)
        elif lower.endswith(".txt"):
            raw = extract_text_from_txt(file_bytes)
        else:
            raise ValueError(f"Unsupported file type: {filename}")
    except Exception as e:
        if isinstance(e, ValueError):
            raise e
        raise ValueError(f"Failed to parse {filename}. The file may be corrupted or malformed. Details: {e}")
    return clean_text(raw)


def clean_text(raw: str) -> str:
    """
    Strip boilerplate noise: extra whitespace, page-break artifacts,
    repeated bullets, control characters. Keeps payload small + clean
    before it ever reaches an LLM (saves context budget + improves matching).
    """
    if not raw:
        return ""
    text = raw.replace("\x00", " ")
    # normalize bullets/symbols to plain dashes
    text = re.sub(r"[•◦▪➤➢●○]", "-", text)
    # collapse multiple blank lines
    text = re.sub(r"\n\s*\n+", "\n", text)
    # collapse runs of spaces/tabs
    text = re.sub(r"[ \t]+", " ", text)
    # strip leading/trailing whitespace per line
    lines = [line.strip() for line in text.split("\n")]
    lines = [line for line in lines if line]
    return "\n".join(lines)


def extract_candidate_name(text: str, fallback: str) -> str:
    """Best-effort guess at candidate name: usually the first non-empty line
    that isn't an email/phone/address and is reasonably short."""
    lines = text.split("\n")[:5]
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if "@" in line or re.search(r"\d{3,}", line):
            continue
        if len(line.split()) <= 5 and len(line) < 60:
            return line
    return fallback


def extract_contact_info(text: str) -> dict:
    """
    Extracts first email and phone number found in the text.
    """
    email_pattern = r"\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\b"
    phone_pattern = r"(?<!\w)(?:\+?\d{1,3}[-.\s]?)?\(?\d{2,4}\)?[-.\s]?\d{2,5}[-.\s]?\d{3,5}\b"
    
    emails = re.findall(email_pattern, text)
    emails = [e for e in emails if "REDACTED" not in e.upper()]
    
    phones = re.findall(phone_pattern, text)
    phones = [p for p in phones if "REDACTED" not in p.upper()]
    
    return {
        "email": emails[0] if emails else "Not Provided",
        "phone": phones[0] if phones else "Not Provided"
    }


def redact_pii(text: str) -> str:
    """
    Redacts emails, phone numbers, and candidate names to support anonymized ranking.
    """
    # 1. Redact Emails
    email_pattern = r"\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\b"
    text = re.sub(email_pattern, "[REDACTED_EMAIL]", text)

    # 2. Redact Phone Numbers
    phone_pattern = r"(?<!\w)(?:\+?\d{1,3}[-.\s]?)?\(?\d{2,4}\)?[-.\s]?\d{2,5}[-.\s]?\d{3,5}\b"
    text = re.sub(phone_pattern, "[REDACTED_PHONE]", text)

    # 3. Redact Candidate Name
    name = extract_candidate_name(text, fallback="")
    if name and name.lower() not in ["resume", "cv", "summary", "experience", "education", "skills"]:
        name_esc = re.escape(name)
        text = re.sub(r"\b" + name_esc + r"\b", "[REDACTED_NAME]", text, flags=re.IGNORECASE)
        # Redact parts of the name
        parts = [p for p in name.split() if len(p) > 2 and p.lower() not in ["and", "the", "for", "with", "from"]]
        for part in parts:
            part_esc = re.escape(part)
            text = re.sub(r"\b" + part_esc + r"\b", "[REDACTED_NAME]", text, flags=re.IGNORECASE)

    return text
